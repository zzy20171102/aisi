"""资料入库：多格式文本抽取 + 分块锚点化（12-Factor #3：上下文按需加载）。"""
from __future__ import annotations

import hashlib
import html as html_mod
import json
import re
from datetime import date, datetime
from pathlib import Path

SUPPORTED = {".md", ".markdown", ".txt", ".text", ".html", ".htm", ".docx", ".pdf"}
MAX_CHUNK_CHARS = 1800


def _extract_plain(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="replace")


def _extract_html(p: Path) -> str:
    raw = p.read_text(encoding="utf-8", errors="replace")
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", "", raw)
    raw = re.sub(r"(?i)<(br|/p|/div|/tr|/li|/h[1-6]|/table)[^>]*>", "\n", raw)
    text = re.sub(r"(?s)<[^>]+>", "", raw)
    return html_mod.unescape(text)


def _extract_docx(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = (para.style.name or "").lower()
        if "heading 1" in style:
            lines.append(f"# {t}")
        elif "heading 2" in style:
            lines.append(f"## {t}")
        elif "heading 3" in style:
            lines.append(f"### {t}")
        else:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(lines)


def _extract_pdf(p: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(p))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        parts.append(f"=== [PDF page {i}] ===")
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def extract(p: Path) -> tuple[str, str]:
    """返回 (文本, 格式)。异常向上抛出由 CLI 转为契约错误。"""
    ext = p.suffix.lower()
    if ext in (".md", ".markdown", ".txt", ".text"):
        return _extract_plain(p), "txt" if ext != ".md" else "md"
    if ext in (".html", ".htm"):
        return _extract_html(p), "html"
    if ext == ".docx":
        return _extract_docx(p), "docx"
    if ext == ".pdf":
        return _extract_pdf(p), "pdf"
    raise ValueError(f"不支持的格式 {ext}（支持：{' '.join(sorted(SUPPORTED))} 或 http 开头的 URL 登记）")


def _is_heading(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^#{1,6}\s+\S", s) or re.match(r"^【.+】$", s)
                or re.match(r"^第.+[章节篇][^\n]{0,30}$", s)
                or (s.endswith(("：", ":")) and len(s) < 40))


def chunk_text(text: str, src_id: str, max_chars: int = MAX_CHUNK_CHARS) -> list[dict]:
    """按空行分块并聚合到 max_chars，每块记录锚点（行号或 PDF 页码）与最近标题。"""
    lines = text.splitlines()
    blocks, cur, start = [], [], 0
    for i, ln in enumerate(lines, 1):
        if ln.strip():
            if not cur:
                start = i
            cur.append((i, ln))
        elif cur:
            blocks.append((start, cur))
            cur = []
    if cur:
        blocks.append((start, cur))

    chunks, heading, page = [], "", None
    buf: list[str] = []
    buf_start, buf_page, buf_heading = None, None, ""
    for bstart, blk in blocks:
        block_first_heading = None
        for ln_no, ln in blk:
            if _is_heading(ln):
                h = ln.strip().lstrip("#").strip()[:60]
                block_first_heading = block_first_heading or h
                heading = h
            m = re.match(r"^=== \[PDF page (\d+)\] ===$", ln.strip())
            if m:
                page = int(m.group(1))
        block_text = "\n".join(x for _, x in blk)
        if buf and sum(len(x) for x in buf) + len(block_text) > max_chars:
            chunks.append(_mk_chunk(src_id, buf, buf_start, buf_page, buf_heading))
            buf, buf_start, buf_page, buf_heading = [], None, None, ""
        if not buf:
            buf_start = bstart
            buf_page = page
            buf_heading = block_first_heading or heading
        buf.append(block_text)
    if buf:
        chunks.append(_mk_chunk(src_id, buf, buf_start, buf_page, buf_heading))
    return chunks


def _mk_chunk(src_id, buf, start, page, heading) -> dict:
    text = "\n\n".join(buf)
    end = start + len("\n".join(buf).splitlines())
    anchor = f"{src_id}#P{page}" if page is not None else f"{src_id}#L{start}-L{end}"
    return {"anchor": anchor, "heading": heading, "page": page, "text": text}


def load_index(ws) -> dict:
    p = ws.base / "sources" / "index.json"
    if p.exists():
        return json.loads(p.read_text(encoding="utf-8"))
    return {"schema": "aisi.sources/1", "sources": []}


def save_index(ws, index: dict) -> None:
    p = ws.base / "sources" / "index.json"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(index, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def next_src_id(ws) -> str:
    ids = {s["id"] for s in load_index(ws)["sources"]}
    for f in (ws.base / "sources").glob("SRC-*.md"):
        ids.add(f.stem.split("-")[0] + "-" + f.stem.split("-")[1])
    n = 0
    for i in ids:
        try:
            n = max(n, int(i.split("-")[1]))
        except (IndexError, ValueError):
            pass
    return f"SRC-{n + 1:03d}"


def ingest_file(ws, file_arg: str, title: str = "") -> dict:
    if file_arg.startswith(("http://", "https://")):
        return _register_url(ws, file_arg, title)
    p = Path(file_arg)
    if not p.exists():
        raise FileNotFoundError(f"资料文件不存在: {p}")
    text, fmt = extract(p)
    src_id = next_src_id(ws)
    chunks = chunk_text(text, src_id)
    title = title or p.stem
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title).strip("-").lower()[:30] or "source"
    out = ws.base / "sources" / f"{src_id}-{slug}.md"
    sha = hashlib.sha256(p.read_bytes()).hexdigest()[:16]
    head = ["---", f"schema: aisi.source/1", f"id: {src_id}", f"title: {title}",
            f"origin: {p.name}", f"format: {fmt}", f"ingested: {date.today().isoformat()}",
            f"chunks: {len(chunks)}", f"sha256: {sha}", "---", ""]
    body = []
    for i, c in enumerate(chunks, 1):
        meta = f"anchor: {c['anchor']}"
        if c["heading"]:
            meta += f" | heading: {c['heading']}"
        body.append(f"<!-- chunk {i} | {meta} -->")
        body.append("")
        body.append(c["text"])
        body.append("")
    out.write_text("\n".join(head + body), encoding="utf-8")
    index = load_index(ws)
    index["sources"] = [s for s in index["sources"] if s["id"] != src_id]
    index["sources"].append({"id": src_id, "title": title, "origin": p.name, "format": fmt,
                              "ingested": date.today().isoformat(), "chunks": len(chunks),
                              "sha256": sha})
    save_index(ws, index)
    return {"id": src_id, "file": str(out), "format": fmt, "chunks": len(chunks), "chars": len(text)}


def _register_url(ws, url: str, title: str) -> dict:
    src_id = next_src_id(ws)
    title = title or url.split("/")[-1][:40] or "web-source"
    out = ws.base / "sources" / f"{src_id}-web.md"
    out.write_text(
        "---\nschema: aisi.source/1\n"
        f"id: {src_id}\ntitle: {title}\norigin: {url}\nformat: url\n"
        f"ingested: {date.today().isoformat()}\nchunks: 0\nsha256: \"\"\n---\n\n"
        "<!-- URL 登记占位：CLI 不联网（12-Factor #11）。请宿主 Agent 抓取页面后，\n"
        "     将正文替换到本文件并按块补 anchor 注释，或直接 ingest 本地保存的文件。 -->\n",
        encoding="utf-8")
    index = load_index(ws)
    index["sources"].append({"id": src_id, "title": title, "origin": url, "format": "url",
                              "ingested": date.today().isoformat(), "chunks": 0, "sha256": ""})
    save_index(ws, index)
    return {"id": src_id, "file": str(out), "format": "url", "chunks": 0, "chars": 0,
            "note": "URL 占位登记，正文待宿主抓取回填"}
