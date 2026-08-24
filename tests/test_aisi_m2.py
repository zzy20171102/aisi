#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AISI M2 资料通道测试：ingest / coverage / research 闭环。"""
import json
import os
import tempfile
import unittest
from pathlib import Path

from aisi import cli
from aisi.ingest import chunk_text
from aisi.validator import load_schema, validate
from tests.test_aisi_m1 import run_cli

SRC_TEXT = """# 人事管理系统设计报告

## 0 概述

本报告详细描述某机关单位人事管理系统的设计与实现，包括研发背景、技术选型、系统流程分析、数据库设计与各功能模块实现。系统采用 B/S 架构与前后端分离模式，后端使用 Spring Boot 框架，前端使用 Vue.js 与 ElementUI 组件库，数据库采用 MySQL 8.0。 """ + "本段落为填充内容，用于验证分块逻辑在超过单块上限时能够正确切分并保留标题锚点。" * 50 + """

## 2.1 访问控制流程

用户登录后获取其所有角色并加载对应权限菜单。对直接输入 URL 的每一个请求，都通过鉴权处理。

## 1.1 可行性分析

系统打包后整体大小只有 38MB，普通单核 CPU、2G RAM 的 ECS 即可满足运行需求。
"""


class M2Tests(unittest.TestCase):
    def setUp(self):
        self.old_cwd = os.getcwd()
        self.tmp = Path(tempfile.mkdtemp(prefix="aisi-m2-"))
        os.chdir(self.tmp)
        run_cli("init", "--id", "hr", "--name", "人力资源管理系统", "--domain", "information-system")
        self.ws = self.tmp / "systems" / "hr"

    def tearDown(self):
        os.chdir(self.old_cwd)
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def write_reqs(self):
        p = self.ws / "views" / "requirements.json"
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(json.dumps({
            "schema": "aisi.requirements/1",
            "requirements": [
                {"id": "REQ-001", "name": "RBAC鉴权", "text": "系统应实现RBAC访问控制鉴权。",
                 "type": "security", "priority": "必须", "verification": "测试",
                 "status": "confirmed", "confidence": "high",
                 "source_refs": ["SRC-001#L1"], "parent": ""},
                {"id": "REQ-002", "name": "性能要求", "text": "系统应响应每一步操作。",
                 "type": "performance", "priority": "应当", "verification": "演示",
                 "status": "proposed", "confidence": "medium",
                 "source_refs": [], "parent": "",
                 "clarifications": [{"question": "响应时间上限？", "answer": ""}]},
            ]}, ensure_ascii=False), encoding="utf-8")

    def test_ingest_md_with_anchors_and_index(self):
        f = self.tmp / "report.md"
        f.write_text(SRC_TEXT, encoding="utf-8")
        rc, out = run_cli("ingest", "--file", str(f), "--title", "人事系统报告")
        self.assertEqual(rc, 0, out)
        self.assertEqual(out["id"], "SRC-001")
        self.assertGreaterEqual(out["chunks"], 1)
        src_file = Path(out["file"])
        text = src_file.read_text(encoding="utf-8")
        self.assertIn("anchor: SRC-001#L", text)
        self.assertIn("heading: 2.1 访问控制流程", text)
        index = json.loads((self.ws / "sources" / "index.json").read_text(encoding="utf-8"))
        self.assertEqual(validate(index, load_schema("sources")), [])
        # 第二次入库编号递增
        rc, out2 = run_cli("ingest", "--file", str(f))
        self.assertEqual(out2["id"], "SRC-002")

    def test_ingest_docx(self):
        from docx import Document
        d = Document()
        d.add_heading("系统需求", 1)
        d.add_paragraph("系统应支持员工资料管理。")
        f = self.tmp / "doc.docx"
        d.save(str(f))
        rc, out = run_cli("ingest", "--file", str(f))
        self.assertEqual(rc, 0, out)
        self.assertEqual(out["format"], "docx")
        self.assertIn("# 系统需求", Path(out["file"]).read_text(encoding="utf-8"))

    def test_ingest_url_placeholder(self):
        rc, out = run_cli("ingest", "--file", "https://example.com/article.html")
        self.assertEqual(rc, 0, out)
        self.assertEqual(out["format"], "url")
        self.assertEqual(out["chunks"], 0)

    def test_ingest_unsupported(self):
        f = self.tmp / "x.xyz"
        f.write_text("data", encoding="utf-8")
        rc, out = run_cli("ingest", "--file", str(f))
        self.assertEqual(rc, 2)
        self.assertEqual(out["errors"][0]["code"], "ingest_failed")

    def test_coverage_gaps_and_resolution(self):
        self.write_reqs()
        rc, out = run_cli("coverage")
        self.assertEqual(rc, 0, out)
        kinds = out["by_kind"]
        self.assertEqual(kinds.get("UNRESOLVED_SOURCE"), 1)
        self.assertEqual(kinds.get("NO_SOURCE"), 1)
        self.assertEqual(kinds.get("OPEN_CLARIFICATION"), 1)
        self.assertEqual(kinds.get("MISSING_MEASURE"), 1)
        gaps = json.loads((self.ws / "research" / "gaps.json").read_text(encoding="utf-8"))
        self.assertEqual(validate(gaps, load_schema("gaps")), [])
        # ingest 后 UNRESOLVED 消解
        f = self.tmp / "report.md"
        f.write_text(SRC_TEXT, encoding="utf-8")
        run_cli("ingest", "--file", str(f))
        rc, out = run_cli("coverage")
        self.assertNotIn("UNRESOLVED_SOURCE", out["by_kind"])

    def test_research_plan_and_ingest(self):
        self.write_reqs()
        run_cli("coverage")
        rc, out = run_cli("research", "plan")
        self.assertEqual(rc, 0, out)
        questions = json.loads((self.ws / "research" / "questions.json").read_text(encoding="utf-8"))
        self.assertEqual(validate(questions, load_schema("research")), [])
        self.assertGreaterEqual(len(questions["questions"]), 1)
        findings = {
            "schema": "aisi.research/1", "topic": "测试调研",
            "questions": [{"id": "Q01", "question": "响应时间上限？", "reason": "test",
                           "status": "answered",
                           "findings": [{"summary": "同类系统页面响应通常要求 <3s",
                                         "source_refs": ["WEB-001"], "confidence": "medium"}]}],
            "sources": [{"id": "WEB-001", "url": "https://example.com", "title": "示例",
                         "accessed": "2026-08-24", "reliability": "medium"}]}
        f = self.tmp / "findings.json"
        f.write_text(json.dumps(findings, ensure_ascii=False), encoding="utf-8")
        rc, out = run_cli("research", "ingest", "--file", str(f), "--to-kb")
        self.assertEqual(rc, 0, out)
        self.assertEqual(out["answered"], 1)
        self.assertEqual(len(out["kb_files"]), 1)
        kb_text = Path(out["kb_files"][0]).read_text(encoding="utf-8")
        self.assertIn("id: KB-", kb_text)
        self.assertIn("<3s", kb_text)
        self.assertIn("kb.py save", out["next_action"])

    def test_coverage_counts_research_web_sources(self):
        """调研 WEB- 来源（findings.json）应计入 coverage 的可解析来源。"""
        self.write_reqs()
        f = self.tmp / "report.md"
        f.write_text(SRC_TEXT, encoding="utf-8")
        run_cli("ingest", "--file", str(f))
        findings = {
            "schema": "aisi.research/1", "topic": "测试调研主题",
            "questions": [{"id": "Q01", "question": "响应时间上限是多少？", "reason": "测试原因",
                           "status": "answered",
                           "findings": [{"summary": "示例调研结论内容", "source_refs": ["WEB-001"]}]}],
            "sources": [{"id": "WEB-001", "url": "https://example.com", "title": "示例来源",
                         "accessed": "2026-08-24", "reliability": "medium"}]}
        fp = self.tmp / "findings.json"
        fp.write_text(json.dumps(findings), encoding="utf-8")
        run_cli("research", "ingest", "--file", str(fp))
        p = self.ws / "views" / "requirements.json"
        data = json.loads(p.read_text(encoding="utf-8"))
        data["requirements"][0]["source_refs"].append("WEB-001")
        p.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        rc, out = run_cli("coverage")
        self.assertEqual(rc, 0, out)
        self.assertNotIn("UNRESOLVED_SOURCE", out["by_kind"])

    def test_chunk_text_pdf_page_anchor(self):
        text = "=== [PDF page 3] ===\n\n某系统性能指标说明文字。\n\n第二段内容。"
        chunks = chunk_text(text, "SRC-001")
        self.assertTrue(all(c["anchor"].startswith("SRC-001#P") for c in chunks))


if __name__ == "__main__":
    unittest.main()
